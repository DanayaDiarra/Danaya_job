"""
applicator/cv_generator.py — Generate tailored .docx + PDF CV per job.
"""
import os
import re
import sqlite3
from pathlib import Path

from dotenv import load_dotenv
from loguru import logger

load_dotenv()

DB_PATH = Path(os.getenv("DB_PATH", "data/jobs.db"))
OUTPUT_DIR = Path("data/applications")

CANDIDATE = {
    "name": "Danaya Diarra",
    "email": "diarradanaya5544@gmail.com",
    "phone": "+7-952-217-0325",
    "location": "Saint Petersburg, Russia",
    "linkedin": "linkedin.com/in/danaya-diarra",
    "github": "github.com/danaya-diarra",
}

EXPERIENCE = [
    {
        "title": "Big Data Analyst Intern",
        "company": "Pulkovo Airport",
        "location": "Saint Petersburg, Russia",
        "dates": "Jul 2025 – Sep 2025",
        "bullets": [
            "Built Power BI dashboards reducing reporting time by 40%",
            "Developed time-series forecasting models using Python/Scikit-learn",
            "Performed large-scale EDA on operational data with Pandas/NumPy",
        ],
    },
    {
        "title": "Market Data Analyst Intern",
        "company": "RostovIT",
        "location": "Rostov-on-Don, Russia",
        "dates": "Jan 2025 – Mar 2025",
        "bullets": [
            "Cleaned and transformed datasets using Python and SQL",
            "Produced market trend and competitive analysis reports",
        ],
    },
    {
        "title": "Supply Chain Engineer",
        "company": "Huawei Technologies",
        "location": "Mali",
        "dates": "2024",
        "bullets": [
            "Automated data entry processes with Python, saving 60% manual effort",
            "Built KPI dashboards for 10+ telecom sites",
            "Optimised last-mile logistics routing",
        ],
    },
    {
        "title": "Co-Founder & CTO",
        "company": "FocusLock",
        "location": "Remote",
        "dates": "2026 – Present",
        "bullets": [
            "Leading product development for a mobile productivity app (React Native)",
            "Conducting user research and UX design with Figma",
            "Implementing freemium business model with gamification mechanics",
        ],
    },
]

EDUCATION = [
    {
        "degree": "MSc Business Analytics & Big Data",
        "school": "GSOM SPbSU",
        "dates": "2024 – 2026 (enrolled)",
        "note": "Thesis: Agentic AI for Predictive Maintenance — PyTorch, LangChain, Transformers",
    },
    {
        "degree": "BSc Supply Chain Management",
        "school": "CUPP-BALLAT",
        "dates": "2021 – 2023",
        "note": "Mali",
    },
]

SKILLS = [
    "Python, SQL, PyTorch, Scikit-learn, Pandas, NumPy",
    "LangChain, FAISS, RAG, Agentic AI, LangGraph",
    "Transformers, LSTM, Deep Learning, XGBoost",
    "Streamlit, Power BI, Tableau, Plotly",
    "Git, Docker (basic), Playwright, Jinja2",
    "React Native, Figma",
    "French (native), English (fluent), Russian (intermediate)",
]


def _slugify(text: str) -> str:
    return re.sub(r"[^\w]+", "_", text.lower()).strip("_")[:40]


def generate_cv(job_id: int, db_path: Path = DB_PATH) -> tuple[Path, Path]:
    """
    Generate tailored CV (.docx + .pdf) for a given job_id.
    Returns (docx_path, pdf_path).
    """
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Load job + tailored summary from DB
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute("""
        SELECT j.title, j.company, s.tailored_cv
        FROM jobs j
        LEFT JOIN scored_jobs s ON s.job_id = j.id
        WHERE j.id = ?
    """, (job_id,))
    row = cur.fetchone()
    conn.close()

    if not row:
        raise ValueError(f"Job {job_id} not found in database")

    title = row["title"] or "role"
    company = row["company"] or "company"
    tailored_summary = row["tailored_cv"] or (
        "Data scientist and ML engineer with 3 years of experience in Python, "
        "PyTorch, LangChain, and agentic AI systems. MSc student at GSOM SPbSU."
    )

    slug = _slugify(f"{company}_{title}")
    docx_path = OUTPUT_DIR / f"cv_{job_id}_{slug}.docx"
    pdf_path = OUTPUT_DIR / f"cv_{job_id}_{slug}.pdf"

    # ── Build .docx ───────────────────────────────────────────────────────
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Narrow margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.9)
            section.right_margin = Inches(0.9)

        # Name header
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_p.add_run(CANDIDATE["name"])
        run.bold = True
        run.font.size = Pt(18)

        # Contact line
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.add_run(
            f"{CANDIDATE['email']} | {CANDIDATE['phone']} | "
            f"{CANDIDATE['location']} | {CANDIDATE['linkedin']}"
        ).font.size = Pt(9)

        doc.add_paragraph()  # spacer

        def section_heading(text):
            p = doc.add_paragraph()
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 100, 160)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            # Horizontal rule via bottom border would need xml, skip for simplicity

        # Professional Summary
        section_heading("Professional Summary")
        doc.add_paragraph(tailored_summary).runs[0].font.size = Pt(10)

        # Experience
        section_heading("Experience")
        for exp in EXPERIENCE:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            r1 = p.add_run(f"{exp['title']} — {exp['company']}")
            r1.bold = True
            r1.font.size = Pt(10)
            p.add_run(f"  |  {exp['location']}  |  {exp['dates']}")

            for bullet in exp["bullets"]:
                b = doc.add_paragraph(style="List Bullet")
                b.add_run(bullet).font.size = Pt(9.5)
                b.paragraph_format.left_indent = Inches(0.25)
                b.paragraph_format.space_after = Pt(1)

        # Education
        section_heading("Education")
        for edu in EDUCATION:
            p = doc.add_paragraph()
            r = p.add_run(f"{edu['degree']} — {edu['school']}")
            r.bold = True
            r.font.size = Pt(10)
            p.add_run(f"  |  {edu['dates']}")
            if edu.get("note"):
                doc.add_paragraph(edu["note"]).runs[0].font.size = Pt(9)

        # Skills
        section_heading("Skills")
        for skill_line in SKILLS:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(skill_line).font.size = Pt(9.5)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(1)

        doc.save(docx_path)
        logger.info(f"CV saved: {docx_path}")

    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        raise

    # ── Convert to PDF ────────────────────────────────────────────────────
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        logger.info(f"PDF saved: {pdf_path}")
    except ImportError:
        logger.warning("docx2pdf not available — trying WeasyPrint HTML fallback")
        _html_to_pdf(tailored_summary, pdf_path)
    except Exception as e:
        logger.warning(f"docx2pdf failed ({e}) — trying WeasyPrint HTML fallback")
        _html_to_pdf(tailored_summary, pdf_path)

    return docx_path, pdf_path


def _html_to_pdf(summary: str, pdf_path: Path) -> None:
    """Fallback: render simple HTML CV to PDF using WeasyPrint."""
    from jinja2 import Template
    import html as html_lib

    tpl_path = Path("templates/cv_template.html")
    if tpl_path.exists():
        template = Template(tpl_path.read_text())
    else:
        template = Template(_FALLBACK_CV_HTML)

    html_str = template.render(
        candidate=CANDIDATE,
        summary=html_lib.escape(summary),
        experience=EXPERIENCE,
        education=EDUCATION,
        skills=SKILLS,
    )

    try:
        from weasyprint import HTML
        HTML(string=html_str).write_pdf(str(pdf_path))
        logger.info(f"PDF (WeasyPrint) saved: {pdf_path}")
    except Exception as e:
        logger.error(f"WeasyPrint PDF generation failed: {e}")


_FALLBACK_CV_HTML = """<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
body{font-family:Arial,sans-serif;font-size:11px;margin:40px;color:#222}
h1{font-size:20px;text-align:center;margin-bottom:4px}
.contact{text-align:center;font-size:9px;color:#555;margin-bottom:16px}
h2{font-size:12px;color:#0064a0;text-transform:uppercase;border-bottom:1px solid #ccc;margin-top:12px}
.job-title{font-weight:bold;font-size:11px}
.meta{font-size:9px;color:#555}
ul{margin:4px 0 8px 16px;padding:0}
li{margin-bottom:2px}
</style></head><body>
<h1>{{ candidate.name }}</h1>
<div class="contact">{{ candidate.email }} | {{ candidate.phone }} | {{ candidate.location }}</div>
<h2>Professional Summary</h2><p>{{ summary }}</p>
<h2>Experience</h2>
{% for exp in experience %}
<div class="job-title">{{ exp.title }} — {{ exp.company }}</div>
<div class="meta">{{ exp.location }} | {{ exp.dates }}</div>
<ul>{% for b in exp.bullets %}<li>{{ b }}</li>{% endfor %}</ul>
{% endfor %}
<h2>Education</h2>
{% for edu in education %}
<div class="job-title">{{ edu.degree }} — {{ edu.school }}</div>
<div class="meta">{{ edu.dates }}{% if edu.note %} | {{ edu.note }}{% endif %}</div>
{% endfor %}
<h2>Skills</h2>
<ul>{% for s in skills %}<li>{{ s }}</li>{% endfor %}</ul>
</body></html>"""
